import os
import io
import json
import zipfile
import tempfile
import subprocess
import base64
from pathlib import Path
from typing import List

import easyocr

from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

ocr_reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)


def _get_file_path(file):
    if isinstance(file, str):
        return file
    if hasattr(file, 'path') and file.path:
        return file.path
    if hasattr(file, 'name') and os.path.exists(file.name):
        return file.name
    return str(file)


def _file_to_data_url(path, ext):
    with open(path, 'rb') as f:
        b64 = base64.b64encode(f.read()).decode()
    mime_map = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    }
    return f"data:{mime_map.get(ext, 'application/octet-stream')};base64,{b64}"


@app.get("/", response_class=HTMLResponse)
async def root():
    return "<h1>Word转换服务运行中</h1><p>API端点: /api/convert_docx_to_pdf, /api/convert_docx_to_images, /api/convert_images_to_docx</p>"


@app.get("/health")
async def health():
    return {"status": "ok"}


@app.post("/api/convert_docx_to_pdf")
async def api_docx_to_pdf(file: UploadFile = File(...)):
    from docx import Document
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.write(await file.read())
    tmp.close()
    try:
        fname = os.path.basename(tmp.name)
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, fname)
            with open(docx_path, 'wb') as f:
                f.write(open(tmp.name, 'rb').read())

            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', tmpdir, docx_path
            ], capture_output=True, timeout=120, text=True)

            if result.returncode != 0:
                return {"error": f"LibreOffice failed: {result.stderr}"}

            pdf_path = os.path.join(tmpdir, Path(docx_path).stem + '.pdf')
            if not os.path.exists(pdf_path):
                return {"error": "PDF conversion failed"}

            return {"data_url": _file_to_data_url(pdf_path, "pdf"),
                    "filename": Path(fname).stem + ".pdf"}
    except Exception as e:
        return {"error": str(e)}
    finally:
        os.unlink(tmp.name)


@app.post("/api/convert_docx_to_images")
async def api_docx_to_images(
    file: UploadFile = File(...),
    format: str = Form("png"),
    dpi: str = Form("300")
):
    orig_name = Path(file.filename).stem
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.write(await file.read())
    tmp.close()
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, file.filename)
            with open(docx_path, 'wb') as f:
                f.write(open(tmp.name, 'rb').read())

            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', tmpdir, docx_path
            ], capture_output=True, timeout=120, text=True)

            if result.returncode != 0:
                return {"error": f"LibreOffice failed: {result.stderr}"}

            pdf_path = os.path.join(tmpdir, Path(docx_path).stem + '.pdf')
            if not os.path.exists(pdf_path):
                return {"error": "PDF conversion failed"}

            from pdf2image import convert_from_path
            images = convert_from_path(pdf_path, dpi=int(dpi))

            images_data = []
            for i, img in enumerate(images):
                buf = io.BytesIO()
                img.save(buf, format=format.upper())
                b64 = base64.b64encode(buf.getvalue()).decode()
                mime = 'image/png' if format == 'png' else 'image/jpeg'
                images_data.append({
                    "dataUrl": f"data:{mime};base64,{b64}",
                    "filename": f"{orig_name}.{format}" if len(images) == 1 else f"{orig_name}_{i + 1}.{format}"
                })

            if len(images_data) == 1:
                return {"images": images_data}

            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                for img_data in images_data:
                    img_bytes = base64.b64decode(img_data["dataUrl"].split(",")[1])
                    zf.writestr(img_data["filename"], img_bytes)
            zip_b64 = base64.b64encode(zip_buf.getvalue()).decode()

            return {"zip_url": f"data:application/zip;base64,{zip_b64}",
                    "filename": f"{orig_name}_images.zip"}
    except Exception as e:
        return {"error": str(e)}
    finally:
        os.unlink(tmp.name)


@app.post("/api/convert_images_to_docx")
async def api_images_to_docx(
    files: List[UploadFile] = File(...),
    layout: str = Form("auto")
):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.oxml.ns import qn
        from PIL import Image

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Noto Serif CJK SC'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
        font.size = Pt(12)

        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

        for i, file in enumerate(files):
            tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            tmp.write(await file.read())
            tmp.close()

            try:
                results = ocr_reader.readtext(tmp.name, detail=1, paragraph=True)

                if results:
                    for item in results:
                        if len(item) == 3:
                            bbox, text, conf = item
                        else:
                            bbox, text = item
                            conf = 1.0

                        text = text.strip()
                        if not text:
                            continue

                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(4)
                        p.paragraph_format.line_spacing = Pt(18)
                        run = p.add_run(text)
                        run.font.name = 'Noto Serif CJK SC'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
                        run.font.size = Pt(12)
                else:
                    doc.add_paragraph(f"[第{i+1}页 - 未能识别文字]")

            finally:
                os.unlink(tmp.name)

        output_path = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc.save(output_path.name)
        output_path.close()

        base_name = Path(files[0].filename).stem if files else "converted"
        return {"data_url": _file_to_data_url(output_path.name, "docx"),
                "filename": base_name + ".docx"}
    except Exception as e:
        return {"error": str(e)}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=7860)
